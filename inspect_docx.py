#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
查看Word文档内容结构
"""

from docx import Document

# 打开文档
doc_path = '/home/admin/.openclaw/workspace-codingagent/projects/剑桥1-3单词.docx'
doc = Document(doc_path)

print("=== 文档段落 ===")
for i, para in enumerate(doc.paragraphs[:50]):  # 先看前50段
    print(f"[{i}] {para.text}")
    # 查看段落样式和字体
    if para.runs:
        for j, run in enumerate(para.runs):
            print(f"  运行[{j}]: 文本='{run.text}', 粗体={run.bold}, 斜体={run.italic}, 颜色={run.font.color.rgb if run.font.color else None}")

print("\n=== 文档表格 ===")
for i, table in enumerate(doc.tables):
    print(f"\n表格[{i}]: {len(table.rows)}行 x {len(table.columns)}列")
    for j, row in enumerate(table.rows[:10]):  # 每个表看前10行
        cells = [cell.text.strip() for cell in row.cells]
        print(f"  行[{j}]: {cells}")
