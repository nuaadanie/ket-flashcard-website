#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
分析文档中颜色的使用情况
"""

from docx import Document

doc_path = '/home/admin/.openclaw/workspace-codingagent/projects/剑桥1-3单词.docx'
doc = Document(doc_path)

# 收集所有颜色
colors = set()
color_examples = {}

# 跳过前面几段，从单词段落开始
for para_idx, para in enumerate(doc.paragraphs[2:50]):  # 看一些有代表性的段落
    if para.text.strip() and para.runs:
        for run in para.runs:
            if run.text.strip():
                color = None
                if run.font.color and run.font.color.rgb:
                    color = str(run.font.color.rgb)
                colors.add(color)
                if color not in color_examples:
                    color_examples[color] = []
                if len(color_examples[color]) < 3:
                    color_examples[color].append(run.text.strip())

print("所有发现的颜色:")
# 将None转换为字符串以便排序
sorted_colors = sorted([str(c) if c is not None else 'None' for c in colors])
for color_str in sorted_colors:
    color = None if color_str == 'None' else color_str
    print(f"\n颜色: {color}")
    if color in color_examples:
        print(f"示例: {', '.join(color_examples[color])}")

# 详细分析几个段落
print("\n\n详细分析段落:")
for para_idx, para in enumerate(doc.paragraphs[2:10]):
    print(f"\n段落 {para_idx+2}: {para.text[:100]}...")
    for run_idx, run in enumerate(para.runs):
        if run.text.strip():
            color = None
            if run.font.color and run.font.color.rgb:
                color = str(run.font.color.rgb)
            print(f"  Run {run_idx}: '{run.text.strip()}' (color: {color})")
